from docx import Document
from docx.shared import Inches

document = Document()
document.add_heading('Python Word Doc')
document.add_paragraph('This is Devnami tutorial')

document.save('Demo1.docx')